from __future__ import annotations

import shutil
import subprocess
from pathlib import Path
from typing import Any

_WINDOWS_SOFFICE_CANDIDATES = [
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
]


def _find_soffice() -> str:
    found = shutil.which("soffice") or shutil.which("soffice.exe")
    if found:
        return found
    for candidate in _WINDOWS_SOFFICE_CANDIDATES:
        if Path(candidate).exists():
            return candidate
    raise RuntimeError(
        "LibreOffice ('soffice') was not found on this system. "
        "Install it (e.g. 'sudo apt install libreoffice' on Linux, "
        "or download LibreOffice on Windows) and ensure it is on PATH."
    )


def convert_to_pdf(source_path: str, output_dir: str | None = None) -> Path:
    source = Path(source_path)
    if not source.is_file():
        raise ValueError(f"Исходный файл не найден: {source_path}")

    target_dir = Path(output_dir) if output_dir else source.parent
    target_dir.mkdir(parents=True, exist_ok=True)

    soffice = _find_soffice()
    result = subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(target_dir),
            str(source),
        ],
        capture_output=True,
        text=True,
        check=False,
    )

    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice conversion failed (exit code {result.returncode}): "
            f"{result.stderr.strip() or result.stdout.strip()}"
        )

    output_path = target_dir / f"{source.stem}.pdf"
    if not output_path.is_file():
        raise RuntimeError(
            f"LibreOffice reported success but no output file was found at {output_path}. "
            f"stdout: {result.stdout.strip()}"
        )
    return output_path


def read_docx_text(path: str) -> str:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed. Install it with: pip install python-docx"
        ) from exc

    source = Path(path)
    if not source.is_file():
        raise ValueError(f"Файл не найден: {path}")

    document = docx.Document(str(source))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def write_docx_text(path: str, text: str) -> Path:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed. Install it with: pip install python-docx"
        ) from exc

    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = docx.Document()
    for line in text.split("\n"):
        document.add_paragraph(line)
    document.save(str(output_path))
    return output_path


def read_xlsx_sheet(path: str, sheet_name: str | None = None) -> list[list[Any]]:
    try:
        import openpyxl
    except ImportError as exc:
        raise RuntimeError(
            "openpyxl is not installed. Install it with: pip install openpyxl"
        ) from exc

    source = Path(path)
    if not source.is_file():
        raise ValueError(f"Файл не найден: {path}")

    workbook = openpyxl.load_workbook(str(source), data_only=True)
    sheet = workbook[sheet_name] if sheet_name else workbook.active
    return [list(row) for row in sheet.iter_rows(values_only=True)]


def write_xlsx_sheet(path: str, rows: list[list[Any]], sheet_name: str = "Sheet1") -> Path:
    try:
        import openpyxl
    except ImportError as exc:
        raise RuntimeError(
            "openpyxl is not installed. Install it with: pip install openpyxl"
        ) from exc

    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = sheet_name
    for row in rows:
        sheet.append(row)
    workbook.save(str(output_path))
    return output_path


class LibreOfficeFileConverter:
    """Adapter satisfying modules.files.ports.FileConverterPort — thin
    delegation to the module-level functions above, which stay as plain
    functions since they have no shared state to encapsulate."""

    convert_to_pdf = staticmethod(convert_to_pdf)
    read_docx_text = staticmethod(read_docx_text)
    write_docx_text = staticmethod(write_docx_text)
    read_xlsx_sheet = staticmethod(read_xlsx_sheet)
    write_xlsx_sheet = staticmethod(write_xlsx_sheet)
